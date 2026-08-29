from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from build_package import (
    ASSET_NAMES,
    FORBIDDEN_PATTERNS,
    FULLWIDTH_ASCII,
    SLOT_RATIOS,
    load_manifest,
    resolve_manifest_path,
    validate_publication_metadata,
)
from user_profile import load_user_profile


CONTENT_INDICES = [3, 4, 6, 7, 8, 10, 11, 13, 14, 16, 17, 19, 20, 22, 23, 24]
FORMULA_RE = re.compile(r"(?:\\frac|\\sum|\\begin\{equation|∑|∫)")


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def main() -> None:
    parser = argparse.ArgumentParser(description="Validate a generated compot-writer package.")
    parser.add_argument("manifest", type=Path)
    parser.add_argument("--output-parent", required=True, type=Path)
    parser.add_argument("--profile-dir", type=Path, help="Override the per-user profile directory")
    parser.add_argument("--fast", action="store_true", help="Validate the folder/DOCX only; skip outer ZIP CRC and content checks")
    args = parser.parse_args()

    manifest_path = args.manifest.expanduser().resolve()
    profile = load_user_profile(args.profile_dir)
    default_author = profile["name"] if profile is not None else None
    manifest, _ = load_manifest(manifest_path, default_author=default_author)
    output_parent = args.output_parent.expanduser().resolve()
    number = int(manifest["number"])
    title = str(manifest["title"])
    author = str(manifest["author"])
    folder_name = f"{number} {title}"
    output_dir = output_parent / folder_name
    docx_path = output_dir / f"[文献速递No.{number}]{title}.docx"
    zip_path = output_parent / f"{folder_name}-{author}.zip"
    source_pdf = resolve_manifest_path(str(manifest["source_pdf"]), manifest_path)
    publish_slot = str(manifest.get("publish_slot") or "待定").strip().replace(":", "：")

    failures: list[str] = []
    warnings: list[str] = []
    if not docx_path.is_file():
        failures.append(f"Missing DOCX: {docx_path}")
    if not args.fast and not zip_path.is_file():
        failures.append(f"Missing ZIP: {zip_path}")
    if failures:
        print(json.dumps({"ok": False, "failures": failures}, ensure_ascii=False, indent=2))
        raise SystemExit(1)

    expected_files = {
        source_pdf.name,
        docx_path.name,
        *ASSET_NAMES,
        f"{author}.png",
        f"{author} {publish_slot}发表.txt",
    }
    actual_files = {path.name for path in output_dir.iterdir() if path.is_file()}
    if actual_files != expected_files:
        failures.append(f"Folder file set mismatch: expected {sorted(expected_files)}, found {sorted(actual_files)}")

    document = Document(docx_path)
    if len(document.paragraphs) != 36:
        failures.append(f"Expected 36 paragraphs, found {len(document.paragraphs)}")
    if len(document.inline_shapes) != 8:
        failures.append(f"Expected 8 inline images, found {len(document.inline_shapes)}")
    if len(document.sections) != 1:
        failures.append(f"Expected one section, found {len(document.sections)}")
    else:
        section = document.sections[0]
        checks = {
            "page width": (section.page_width.inches, 8.2677),
            "page height": (section.page_height.inches, 11.6929),
            "top margin": (section.top_margin.inches, 1.0),
            "bottom margin": (section.bottom_margin.inches, 1.0),
            "left margin": (section.left_margin.inches, 1.25),
            "right margin": (section.right_margin.inches, 1.25),
        }
        for label, (actual, expected) in checks.items():
            if abs(actual - expected) > 0.03:
                failures.append(f"{label} is {actual:.3f}, expected {expected:.3f}")

    expected_title = f"[文献速递No.{number}]{title}"
    if document.paragraphs[0].text != expected_title:
        failures.append("Title text does not match the naming contract")
    if document.paragraphs[1].text != f"撰稿人：{author}":
        failures.append("Opening writer line is missing or misplaced")
    if document.paragraphs[33].paragraph_format.page_break_before is not True:
        failures.append("Final writer card must start on a new page")
    for run in document.paragraphs[0].runs:
        if not run.text:
            continue
        rpr = run._element.rPr
        fonts = rpr.rFonts if rpr is not None else None
        east_asia = fonts.get(qn("w:eastAsia")) if fonts is not None else None
        if east_asia != "黑体" or run.bold is not True:
            failures.append(f"Title run is not bold SimHei for Chinese text: {run.text!r}")
        if run.font.size is None or abs(run.font.size.pt - 15.0) > 0.1:
            failures.append(f"Title run is not 15 pt: {run.text!r}")

    article = "\n".join(document.paragraphs[index].text for index in CONTENT_INDICES)
    forbidden = [label for label, pattern in FORBIDDEN_PATTERNS.items() if pattern.search(article)]
    if forbidden:
        failures.append(f"Forbidden self-attribution remains: {forbidden}")
    remaining = FULLWIDTH_ASCII.findall(article)
    if remaining:
        failures.append(f"Fullwidth parentheses remain around ASCII-only content: {remaining}")
    if FORMULA_RE.search(article):
        failures.append("Formula markers remain in the authored article")
    try:
        validate_publication_metadata(
            f"{document.paragraphs[3].text}\n{document.paragraphs[4].text}"
        )
    except (IndexError, ValueError) as exc:
        failures.append(f"Publication metadata format is invalid: {exc}")
    body_chars = sum(len(document.paragraphs[index].text) for index in CONTENT_INDICES)
    if not 2200 <= body_chars <= 3600:
        warnings.append(f"Authored body length is {body_chars} characters; reference target is roughly 2500–3200")

    for name, expected_ratio in SLOT_RATIOS.items():
        path = output_dir / name
        if not path.is_file():
            continue
        with Image.open(path) as image:
            ratio = image.width / image.height
        if abs(ratio / expected_ratio - 1.0) > 0.03:
            failures.append(f"{name} ratio {ratio:.3f} differs from slot {expected_ratio:.3f}")

    with zipfile.ZipFile(docx_path) as package:
        if not args.fast:
            bad = package.testzip()
            if bad:
                failures.append(f"Corrupt DOCX package member: {bad}")
        media = [name for name in package.namelist() if name.startswith("word/media/")]
        if len(media) != 8:
            failures.append(f"Expected 8 DOCX media parts, found {len(media)}")

    if not args.fast:
        with zipfile.ZipFile(zip_path) as archive:
            root = f"{folder_name}/"
            expected_entries = {root, *(f"{root}{name}" for name in expected_files)}
            actual_entries = set(archive.namelist())
            if actual_entries != expected_entries:
                failures.append("Outer ZIP entry set does not match the final folder")
            # Reading every member once verifies CRC. Hash the embedded DOCX in
            # that same pass instead of calling testzip() and reading it twice.
            embedded_docx_hash = None
            for name in sorted(actual_entries):
                if name.endswith("/"):
                    continue
                try:
                    data = archive.read(name)
                except (zipfile.BadZipFile, RuntimeError) as exc:
                    failures.append(f"Corrupt outer ZIP member {name}: {exc}")
                    break
                if name == f"{root}{docx_path.name}":
                    embedded_docx_hash = sha256(data)
            if embedded_docx_hash is not None and embedded_docx_hash != sha256(docx_path.read_bytes()):
                failures.append("The DOCX inside the ZIP is not the final DOCX")

    report = {
        "ok": not failures,
        "folder": str(output_dir),
        "docx": str(docx_path),
        "zip": str(zip_path),
        "mode": "fast" if args.fast else "final",
        "body_characters": body_chars,
        "failures": failures,
        "warnings": warnings,
        "note": "Structural checks do not replace rendering and visual inspection of every page.",
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if failures:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
